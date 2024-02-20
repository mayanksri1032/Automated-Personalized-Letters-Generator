import pandas as pd
from docx import Document

# Load Excel sheet
excel_file_path = 'your_excel_file.xlsx'
df = pd.read_excel(excel_file_path)

for index, row in df.iterrows():
    person_name = row['Name']
    address = row['Address']
    account_number = row['A/C Number']
    amount = row['Amount']

    # Create a new Word document
    doc = Document()

    # Add content to the Word document
    doc.add_heading('Letter', level=1)
    doc.add_paragraph(f'Dear {person_name},')
    doc.add_paragraph(f'Address: {address}')
    doc.add_paragraph(f'A/C Number: {account_number}')
    doc.add_paragraph(f'Amount: {amount}')

    # Add more content or customize the letter as needed

    # Save the Word document
    word_file_path = f'{person_name}_letter.docx'
    doc.save(word_file_path)

    print(f'Letter for {person_name} saved as {word_file_path}')
